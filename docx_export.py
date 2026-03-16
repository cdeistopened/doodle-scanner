"""Convert OCR markdown output to .docx format."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


def markdown_to_docx(md_text: str, output_path: str):
    """Convert markdown text to a Word document.

    Handles headings (#, ##, ###), bold (**text**), italic (*text*),
    bullet lists (- item), numbered lists (1. item), and paragraphs.
    """
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = md_text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith('### '):
            p = doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith('## '):
            p = doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        # Horizontal rule
        if stripped in ('---', '***', '___'):
            doc.add_paragraph('_' * 40)
            i += 1
            continue

        # Bullet list
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, text)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if num_match:
            text = num_match.group(2)
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, text)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_formatted_text(p, stripped)
        i += 1

    doc.save(output_path)


def _add_formatted_text(paragraph, text: str):
    """Add text with bold and italic formatting to a paragraph."""
    # Pattern matches **bold**, *italic*, and plain text segments
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))'

    for match in re.finditer(pattern, text):
        if match.group(2):  # Bold
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # Italic
            run = paragraph.add_run(match.group(3))
            run.italic = True
        elif match.group(4):  # Plain
            paragraph.add_run(match.group(4))
